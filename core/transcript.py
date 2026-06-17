from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Formatting helpers ────────────────────────────────────────────────────────

def format_timestamp(seconds: float) -> str:
    """Convert float seconds to HH:MM:SS.mmm string (plain transcript style)."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:06.3f}"


def format_srt_timestamp(seconds: float) -> str:
    """Convert float seconds to SRT-style HH:MM:SS,mmm string."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int(round((seconds - int(seconds)) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def format_segments(segments: list[dict]) -> str:
    """
    Build a timestamped transcript string from Groq segment dicts.
    Format: [HH:MM:SS.mmm] text
    """
    lines = []
    for seg in segments:
        ts = format_timestamp(seg["start"])
        lines.append(f"[{ts}] {seg['text'].strip()}")
    return "\n".join(lines)


def format_srt(segments: list[dict]) -> str:
    """
    Build a standard .srt subtitle file from Groq segment dicts.
    Format:
        1
        00:00:00,000 --> 00:00:03,500
        Segment text here.

        2
        ...
    """
    blocks = []
    for i, seg in enumerate(segments, start=1):
        text = seg["text"].strip()
        if not text:
            continue
        start = format_srt_timestamp(seg["start"])
        end = format_srt_timestamp(seg["end"])
        blocks.append(f"{i}\n{start} --> {end}\n{text}")
    return "\n\n".join(blocks) + "\n"


# ── Paragraph segmentation ───────────────────────────────────────────────────

def segment_by_silence(segments: list[dict], threshold: float = 2.0) -> str:
    """
    Join segment text into a plain transcript, inserting a blank line
    (paragraph break) wherever the gap between segments is >= threshold seconds.

    This is step 1 of 2 — the Gemini formatting pass cleans it up after.
    """
    if not segments:
        return ""

    paragraphs = []
    current: list[str] = []

    for i, seg in enumerate(segments):
        text = seg["text"].strip()
        if not text:
            continue

        if i > 0:
            gap = seg["start"] - segments[i - 1]["end"]
            if gap >= threshold:
                if current:
                    paragraphs.append(" ".join(current))
                    current = []

        current.append(text)

    if current:
        paragraphs.append(" ".join(current))

    return "\n\n".join(paragraphs)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_transcript_docx(title: str, plain_text: str, output_path: Path) -> Path:
    """
    Build a .docx transcript with:
      - Title
      - A horizontal divider line
      - The formatted paragraphs
    """
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Divider — bottom border on an empty paragraph
    divider = doc.add_paragraph()
    p_pr = divider._p.get_or_add_pPr()
    p_borders = p_pr.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr", {}
    )
    bottom = p_borders.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom",
        {
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "single",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz": "6",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space": "1",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color": "888888",
        },
    )
    p_borders.append(bottom)
    p_pr.append(p_borders)

    # Paragraphs
    for para_text in plain_text.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(12)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ── Output writers ────────────────────────────────────────────────────────────

def write_transcript_files(
    segments: list[dict],
    plain_text: str,
    output_dir: Path,
    episode_slug: str,
    episode_title: str | None = None,
) -> dict[str, Path]:
    """
    Write transcript variants to disk:
      - timestamped -> .srt
      - plain       -> .docx
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    paths = {
        "timestamped": output_dir / f"{episode_slug}_timestamped.srt",
        "plain": output_dir / f"{episode_slug}_plain.docx",
    }

    paths["timestamped"].write_text(format_srt(segments), encoding="utf-8")
    build_transcript_docx(episode_title or episode_slug, plain_text, paths["plain"])

    return paths